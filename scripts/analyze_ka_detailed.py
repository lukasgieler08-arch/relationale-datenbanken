#!/usr/bin/env python3
"""
Erweiterte Analyse der Klassenarbeit mit Tabellen-Unterstützung
"""

from docx import Document
from docx.table import Table
import json
import re

def analyze_docx_detailed(filepath):
    """Detaillierte Analyse mit Tabellen"""
    doc = Document(filepath)
    
    content = {
        "title": "",
        "aufgaben": [],
        "tabellen": [],
        "all_text": [],
    }
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            content["all_text"].append(text)
            if para.style.name.startswith('Heading'):
                content["title"] = text
    
    # Tabellen extrahieren
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        content["tabellen"].append(table_data)
    
    return content

def main():
    ka_file = "/workspaces/edu-code-course-rdb/uploads/klassenarbeiten-und-unterrichtsmaterialien/KA02_INF_BG12_2024_2025_CJ_LSG_VERSION3.docx"
    
    print("🔍 Detaillierte Analyse der Klassenarbeit...\n")
    result = analyze_docx_detailed(ka_file)
    
    print(f"📄 Titel: {result['title']}\n")
    
    print(f"📊 Gefundene Tabellen: {len(result['tabellen'])}\n")
    for i, table in enumerate(result['tabellen'], 1):
        print(f"\nTabelle {i}:")
        for row in table[:5]:  # Erste 5 Zeilen
            print(f"  {row}")
        if len(table) > 5:
            print(f"  ... ({len(table) - 5} weitere Zeilen)")
    
    print(f"\n📝 Text-Zusammenfassung ({len(result['all_text'])} Zeilen):")
    for i, line in enumerate(result['all_text'][:30]):
        print(f"  {i+1}. {line[:80]}")
    
    if len(result['all_text']) > 30:
        print(f"  ... ({len(result['all_text']) - 30} weitere Zeilen)")
    
    # Speichere detaillierte Analyse
    with open("/tmp/ka_detailed.json", "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)
    
    print(f"\n✅ Vollständige Analyse gespeichert in: /tmp/ka_detailed.json")

if __name__ == "__main__":
    main()
